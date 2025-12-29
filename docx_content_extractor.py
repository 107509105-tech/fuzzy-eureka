"""
DOCX 內容提取器
支援提取文字、圖片、表格，並將頁面保存為圖片
"""

import os
from typing import List, Optional
from dataclasses import dataclass
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from PIL import Image
import io
import base64
from docx.document import Document as DocumentType


@dataclass
class DocxPageContent:
    """DOCX 頁面內容結構"""
    page_num: int
    page_type: str  # 'text_heavy', 'table_heavy', 'image_heavy', 'mixed'
    text: str
    tables: List[List[List[str]]]  # 表格數據
    images_count: int
    image_path: Optional[str]  # 保存的頁面圖片路徑
    text_coverage: float  # 0-1，文字內容佔比


class DocxContentExtractor:
    """DOCX 內容提取器"""

    PAGE_TYPE_TEXT_HEAVY = "text_heavy"
    PAGE_TYPE_TABLE_HEAVY = "table_heavy"
    PAGE_TYPE_IMAGE_HEAVY = "image_heavy"
    PAGE_TYPE_MIXED = "mixed"

    def __init__(self, output_dir: str = "static/pages"):
        """
        初始化

        Args:
            output_dir: 圖片輸出目錄
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def extract_images_from_docx(self, doc: DocumentType, doc_name: str) -> List[str]:
        """
        從 DOCX 中提取並保存所有圖片

        Args:
            doc: Document 對象
            doc_name: 文件名（不含副檔名）

        Returns:
            List[str]: 圖片路徑列表
        """
        image_paths = []
        image_count = 0

        # 遍歷文件中的所有關係
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_count += 1
                    image_data = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_data))

                    # 保存圖片
                    image_filename = f"{doc_name}_image_{image_count}.png"
                    image_path = os.path.join(self.output_dir, image_filename)
                    image.save(image_path, "PNG")
                    image_paths.append(image_path)
                except Exception as e:
                    print(f"提取圖片 {image_count} 失敗: {e}")
                    continue

        return image_paths

    def extract_tables(self, doc: DocumentType) -> List[List[List[str]]]:
        """
        提取文件中的所有表格

        Args:
            doc: Document 對象

        Returns:
            List[List[List[str]]]: 表格列表，每個表格是一個二維數組
        """
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            if table_data:  # 只添加非空表格
                tables.append(table_data)
        return tables

    def extract_text(self, doc: DocumentType) -> str:
        """
        提取文件中的所有文字（不包括表格）

        Args:
            doc: Document 對象

        Returns:
            str: 提取的文字內容
        """
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)

    def determine_page_type(self, text_length: int, tables_count: int, images_count: int) -> tuple[str, float]:
        """
        判斷頁面類型

        Args:
            text_length: 文字長度
            tables_count: 表格數量
            images_count: 圖片數量

        Returns:
            tuple[str, float]: (頁面類型, 文字覆蓋率)
        """
        total_content = text_length + tables_count * 100 + images_count * 200

        if total_content == 0:
            return self.PAGE_TYPE_TEXT_HEAVY, 0.0

        text_coverage = text_length / total_content if total_content > 0 else 0

        # 判斷頁面類型
        if images_count > 3 or (images_count > 0 and text_coverage < 0.2):
            return self.PAGE_TYPE_IMAGE_HEAVY, text_coverage
        elif tables_count > 2 or (tables_count > 0 and text_coverage < 0.4):
            return self.PAGE_TYPE_TABLE_HEAVY, text_coverage
        elif text_coverage > 0.8:
            return self.PAGE_TYPE_TEXT_HEAVY, text_coverage
        else:
            return self.PAGE_TYPE_MIXED, text_coverage

    def convert_docx_to_image(self, docx_path: str, doc_name: str) -> Optional[str]:
        """
        將 DOCX 轉換為圖片（使用 LibreOffice 或其他工具）
        注意：這需要系統安裝 LibreOffice

        Args:
            docx_path: DOCX 檔案路徑
            doc_name: 文件名（不含副檔名）

        Returns:
            Optional[str]: 圖片路徑，如果轉換失敗則返回 None
        """
        try:
            import subprocess

            # 使用 LibreOffice 轉換為 PDF
            pdf_path = os.path.join(self.output_dir, f"{doc_name}.pdf")
            subprocess.run([
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', self.output_dir,
                docx_path
            ], check=True, capture_output=True)

            # 使用 PyMuPDF 將 PDF 轉換為圖片
            if os.path.exists(pdf_path):
                import fitz
                doc = fitz.open(pdf_path)
                page = doc[0]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))

                image_filename = f"{doc_name}_preview.png"
                image_path = os.path.join(self.output_dir, image_filename)
                pix.save(image_path)

                doc.close()
                os.remove(pdf_path)  # 刪除臨時 PDF

                return image_path
        except Exception as e:
            print(f"DOCX 轉圖片失敗: {e}")

        return None

    def extract(self, docx_path: str, doc_name: Optional[str] = None) -> List[DocxPageContent]:
        """
        提取 DOCX 檔案的所有內容

        Args:
            docx_path: DOCX 檔案路徑
            doc_name: 文件名（不含副檔名），如果為 None 則從路徑中提取

        Returns:
            List[DocxPageContent]: 提取的內容列表（DOCX 視為單頁）
        """
        if doc_name is None:
            doc_name = os.path.splitext(os.path.basename(docx_path))[0]

        try:
            doc = Document(docx_path)

            # 提取內容
            text = self.extract_text(doc)
            tables = self.extract_tables(doc)
            image_paths = self.extract_images_from_docx(doc, doc_name)

            # 判斷頁面類型
            page_type, text_coverage = self.determine_page_type(
                len(text),
                len(tables),
                len(image_paths)
            )

            # 嘗試生成預覽圖片
            preview_image = self.convert_docx_to_image(docx_path, doc_name)

            # 組合表格內容到文字中
            if tables:
                text += "\n\n=== 表格內容 ===\n"
                for idx, table in enumerate(tables):
                    text += f"\n表格 {idx + 1}:\n"
                    for row in table:
                        text += " | ".join(row) + "\n"

            # 建立 DocxPageContent 對象（DOCX 視為單頁，頁碼為 1）
            content = DocxPageContent(
                page_num=1,
                page_type=page_type,
                text=text,
                tables=tables,
                images_count=len(image_paths),
                image_path=preview_image,
                text_coverage=text_coverage
            )

            return [content]

        except Exception as e:
            print(f"提取 DOCX 失敗: {e}")
            raise

    def extract_all_images(self, docx_path: str, doc_name: Optional[str] = None) -> List[str]:
        """
        單獨提取所有圖片（不提取文字和表格）

        Args:
            docx_path: DOCX 檔案路徑
            doc_name: 文件名（不含副檔名）

        Returns:
            List[str]: 圖片路徑列表
        """
        if doc_name is None:
            doc_name = os.path.splitext(os.path.basename(docx_path))[0]

        doc = Document(docx_path)
        return self.extract_images_from_docx(doc, doc_name)


if __name__ == "__main__":
    # 測試代碼
    extractor = DocxContentExtractor()

    # 測試提取
    test_docx = "test.docx"  # 替換為實際的 DOCX 檔案
    if os.path.exists(test_docx):
        contents = extractor.extract(test_docx)
        for content in contents:
            print(f"頁碼: {content.page_num}")
            print(f"類型: {content.page_type}")
            print(f"文字長度: {len(content.text)}")
            print(f"表格數量: {len(content.tables)}")
            print(f"圖片數量: {content.images_count}")
            print(f"圖片路徑: {content.image_path}")
            print("-" * 50)
